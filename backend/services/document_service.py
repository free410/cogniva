from typing import Optional, List, Dict, Any
import os
import uuid
from datetime import datetime
from pathlib import Path
import aiofiles
from sqlalchemy.orm import Session
from fastapi import UploadFile
import pypdf
import docx
import pandas as pd
from models import Document, Chunk, Vector
from core.config import settings
from services.rag_service import embedding_service
from services.advanced_chunker import AdvancedChunker


class DocumentService:
    """文档处理服务"""

    def __init__(self, db: Session):
        self.db = db
        self.upload_dir = Path(settings.UPLOAD_DIR)
        self.upload_dir.mkdir(parents=True, exist_ok=True)

    def _get_file_extension(self, filename: str) -> str:
        return Path(filename).suffix.lower().lstrip(".")

    async def save_file(self, file: UploadFile) -> Dict[str, Any]:
        """保存上传的文件"""
        file_id = str(uuid.uuid4())
        extension = self._get_file_extension(file.filename)
        stored_filename = f"{file_id}.{extension}"
        file_path = self.upload_dir / stored_filename

        content = await file.read()
        async with aiofiles.open(file_path, "wb") as f:
            await f.write(content)

        return {
            "file_id": file_id,
            "filename": file.filename,
            "file_path": str(file_path),
            "file_type": extension,
            "file_size": len(content)
        }

    async def extract_text(self, file_path: str, file_type: str) -> str:
        """从不同格式文件中提取文本"""
        if file_type == "pdf":
            return self._extract_pdf(file_path)
        elif file_type in ["docx", "doc"]:
            return self._extract_docx(file_path)
        elif file_type in ["txt", "md"]:
            return self._extract_txt(file_path)
        elif file_type in ["xlsx", "xls", "csv"]:
            return self._extract_excel(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def _extract_pdf(self, file_path: str) -> str:
        text_parts = []
        with open(file_path, "rb") as f:
            reader = pypdf.PdfReader(f)
            for page in reader.pages:
                text_parts.append(page.extract_text())
        return "\n".join(text_parts)

    def _extract_docx(self, file_path: str) -> str:
        doc = docx.Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs])

    def _extract_txt(self, file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    def _extract_excel(self, file_path: str) -> str:
        if file_path.endswith(".csv"):
            df = pd.read_csv(file_path)
        else:
            df = pd.read_excel(file_path)
        return df.to_string()

    async def process_document(
        self,
        user_id: str,
        file: UploadFile,
        chunk_size: int = None,
        overlap: int = None,
        chunk_strategy: str = None
    ) -> Document:
        """处理上传的文档"""
        if chunk_size is None:
            chunk_size = settings.CHUNK_SIZE
        if overlap is None:
            overlap = settings.CHUNK_OVERLAP
        if chunk_strategy is None:
            chunk_strategy = getattr(settings, 'CHUNK_STRATEGY', 'recursive')

        # 保存文件
        file_info = await self.save_file(file)

        # 创建文档记录
        document = Document(
            user_id=user_id,
            filename=file_info["filename"],
            file_type=file_info["file_type"],
            file_size=file_info["file_size"],
            file_path=file_info["file_path"],
            status="processing",
            extra_metadata={"chunk_strategy": chunk_strategy}
        )
        self.db.add(document)
        self.db.commit()

        try:
            # 提取文本
            text = await self.extract_text(file_info["file_path"], file_info["file_type"])

            # 分块（使用高级策略）
            chunks = self._chunk_text(text, chunk_size, overlap, chunk_strategy)

            # 生成嵌入并保存
            for i, chunk_text in enumerate(chunks):
                chunk = Chunk(
                    document_id=document.id,
                    content=chunk_text,
                    chunk_index=i,
                    token_count=len(chunk_text) // 4,  # 粗略估计
                    extra_metadata={
                        "chunk_strategy": chunk_strategy,
                        "char_count": len(chunk_text),
                        "word_count": len(chunk_text.split())
                    }
                )
                self.db.add(chunk)
                self.db.flush()

                # 生成向量
                embedding = embedding_service.embed([chunk_text])[0]
                vector = Vector(
                    chunk_id=chunk.id,
                    embedding=embedding
                )
                self.db.add(vector)

            document.status = "completed"
            document.extra_metadata = {
                "chunk_strategy": chunk_strategy,
                "chunk_count": len(chunks),
                "total_chars": len(text),
                "avg_chunk_size": sum(len(c) for c in chunks) / len(chunks) if chunks else 0
            }
            self.db.commit()

        except Exception as e:
            document.status = "failed"
            document.extra_metadata = {"error": str(e), "chunk_strategy": chunk_strategy}
            self.db.commit()
            raise

        return document

    def _chunk_text(
        self,
        text: str,
        chunk_size: int = None,
        overlap: int = None,
        strategy: str = None
    ) -> List[str]:
        """
        文本分块 - 支持多种高级切块策略

        Args:
            text: 输入文本
            chunk_size: 块大小
            overlap: 重叠大小
            strategy: 切块策略
                - "smart_hybrid": 智能混合（推荐）
                - "intent": 意图感知切块
                - "entity": 实体感知切块
                - "deep_semantic": 深度语义切块
                - "recursive": 递归字符切块

        Returns:
            切块列表
        """
        if chunk_size is None:
            chunk_size = settings.CHUNK_SIZE
        if overlap is None:
            overlap = settings.CHUNK_OVERLAP
        if strategy is None:
            strategy = getattr(settings, 'CHUNK_STRATEGY', 'smart_hybrid')

        # 预处理文本
        text = self._preprocess_text(text)
        if not text.strip():
            return []

        # 使用高级切块器
        chunker = AdvancedChunker(
            chunk_size=chunk_size,
            chunk_overlap=overlap,
            min_chunk_size=200
        )

        try:
            chunks = chunker.chunk(text, strategy=strategy)
            # 后处理：确保块不为空，且大小合理
            return [c.strip() for c in chunks if c.strip() and len(c.strip()) >= 80]
        except Exception as e:
            print(f"高级切块失败 ({strategy}): {e}, 回退到智能混合切块")
            # 回退到默认策略
            return chunker.chunk(text, strategy="smart_hybrid")

    def _preprocess_text(self, text: str) -> str:
        """文本预处理"""
        # 统一换行符
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        # 去除多余的空格
        text = ' '.join(text.split())
        return text

    def _split_paragraphs(self, text: str) -> List[str]:
        """按段落分割文本"""
        # 按多个空行分割段落
        paragraphs = []
        current_para = []

        for line in text.split('\n'):
            line = line.strip()
            if not line:
                if current_para:
                    paragraphs.append(' '.join(current_para))
                    current_para = []
            else:
                current_para.append(line)

        if current_para:
            paragraphs.append(' '.join(current_para))

        return [p for p in paragraphs if p]

    def _split_long_paragraph(self, paragraph: str, chunk_size: int, overlap: int) -> List[str]:
        """拆分过长的段落 - 按句子边界拆分"""
        sentences = self._split_sentences(paragraph)

        sub_chunks = []
        current_chunk = []
        current_length = 0

        for sentence in sentences:
            sent_len = len(sentence)

            # 如果单个句子就超过块大小的一半，单独成块
            if sent_len > chunk_size * 0.6 and current_length == 0:
                sub_chunks.append(sentence)
                continue

            if current_length + sent_len > chunk_size and current_chunk:
                sub_chunks.append("".join(current_chunk))
                # 保留重叠部分：最后几个完整句子
                current_chunk = self._get_overlap_chunks(current_chunk, overlap)
                current_length = sum(len(c) for c in current_chunk)

            current_chunk.append(sentence)
            current_length += sent_len

        if current_chunk:
            sub_chunks.append("".join(current_chunk))

        return sub_chunks

    def _split_sentences(self, text: str) -> List[str]:
        """按句子分割文本 - 增强版"""
        # 中文句子分割符
        sentence_endings = ['。', '！', '？', '；', '…', '，']
        sentences = []
        current = []

        for char in text:
            current.append(char)
            if char in sentence_endings:
                sentences.append(''.join(current).strip())
                current = []

        if current:
            sentences.append(''.join(current).strip())

        # 过滤空句子
        return [s for s in sentences if s]

    def _get_overlap_chunks(self, chunks: List[str], overlap: int) -> List[str]:
        """获取重叠部分"""
        if not chunks:
            return []

        # 合并所有chunk
        full_text = " ".join(chunks)

        # 取最后 overlap 个字符
        if len(full_text) <= overlap:
            return [full_text]

        overlap_text = full_text[-overlap:]

        # 尝试在句子边界处分割
        for sep in ['。', '！', '？', '. ', '\n']:
            pos = overlap_text.rfind(sep)
            if pos > 0:
                overlap_text = overlap_text[pos+1:].strip()
                break

        return [overlap_text] if overlap_text else []

    async def get_documents(self, user_id: str) -> List[Document]:
        """获取用户的所有文档"""
        return self.db.query(Document).filter(
            Document.user_id == user_id
        ).order_by(Document.created_at.desc()).all()

    async def delete_document(self, document_id: str, user_id: str) -> bool:
        """删除文档及其关联数据"""
        document = self.db.query(Document).filter(
            Document.id == document_id,
            Document.user_id == user_id
        ).first()

        if not document:
            return False

        # 删除文件
        try:
            if os.path.exists(document.file_path):
                os.remove(document.file_path)
        except OSError:
            pass

        # 手动删除关联的 chunks 和 vectors（确保正确顺序）
        from models import Chunk, Vector, Citation

        # 先删除 Citation（引用）
        chunk_ids = [c.id for c in document.chunks]
        if chunk_ids:
            self.db.query(Citation).filter(Citation.chunk_id.in_(chunk_ids)).delete(synchronize_session=False)

        # 删除 vectors
        self.db.query(Vector).filter(Vector.chunk_id.in_(chunk_ids)).delete(synchronize_session=False)

        # 删除 chunks
        self.db.query(Chunk).filter(Chunk.document_id == document_id).delete(synchronize_session=False)

        # 删除文档
        self.db.delete(document)
        self.db.commit()
        return True


document_service = DocumentService